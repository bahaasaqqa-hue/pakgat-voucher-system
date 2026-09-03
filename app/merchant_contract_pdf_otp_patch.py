"""Native DOCX merchant contract renderer with OTP acceptance."""
from __future__ import annotations

import base64
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from app import merchant_contract_pdf as contract_pdf

NAVY = "123D80"
LIGHT = "F7F9FC"
BORDER = "CDD9EA"
TEXT = "172B4D"
MUTED = "526783"
WHITE = "FFFFFF"

ACTIVATION_COPY = (
    "نجاح OTP يثبت موافقة التاجر على الاتفاقية ولا يفعّل الحساب تلقائياً. "
    "يصبح الحساب Active فقط بعد الموافقة النهائية من Pakgat على طلب التسجيل."
)

CLAUSES = (
    (1, "التزامات الطرف الأول (Pakgat)", "تتولى Pakgat تصميم ونشر وتسويق العروض عبر المنصة، وإصدار القسائم الإلكترونية وتوفير نظام التحقق منها، وتحصيل قيمة الطلبات وتحويل صافي مستحقات التاجر وفق البيانات المالية المعتمدة، مع تقديم دعم المنصة. ولا تضمن Pakgat عدداً محدداً من المبيعات أو العملاء."),
    (2, "التزامات الطرف الثاني (التاجر)", "يلتزم التاجر بتقديم الخدمة أو المنتج بالمواصفات والسعر والشروط المعتمدة، وقبول القسائم الصالحة دون رسوم غير معلنة، وضمان صحة بياناته وأسعاره وتراخيصه والمحافظة على وسائل التحقق، ويتحمل المسؤولية النظامية عن جودة ما يقدمه وأي مطالبات ناشئة عنه."),
    (3, "القسائم الإلكترونية", "تصدر لكل عملية شراء قسيمة برمز تحقق فريد تستخدم لمرة واحدة، وتعد مستخدمة بعد اعتمادها عبر النظام، ولا يجوز إعادة استخدامها أو قبولها بعد انتهاء صلاحيتها إلا باستثناء مكتوب من Pakgat لحماية حقوق العميل."),
    (4, "التسوية المالية", "تحول Pakgat صافي مستحقات التاجر إلى الآيبان المسجل وفق دورة التسوية المعتمدة بعد خصم المبالغ المستردة والإلغاءات والأخطاء المحاسبية والرسوم المتفق عليها. ويعد كشف التسوية أساساً للمراجعة، ويكون الاعتراض خلال 7 أيام عمل من إرساله."),
    (5, "الإلغاء والاسترجاع وحقوق العملاء", "تخضع طلبات الإلغاء والاستبدال والاسترجاع لسياسات Pakgat المعلنة وبما يتفق مع الأنظمة. وإذا ألغي الطلب أو تعذر على التاجر تقديم الخدمة أو المنتج لسبب يعود إليه، يلتزم بمعالجة حقوق العملاء وإعادة المبالغ المستحقة أو تنفيذ البديل الذي تعتمده Pakgat."),
    (6, "التسويق والملكية الفكرية", "تبقى حقوق الملكية الفكرية والعلامات التجارية لكل طرف ملكاً له. ويمنح التاجر Pakgat ترخيصاً غير حصري ومجانياً طوال مدة التعاون لاستخدام اسمه وشعاره وصوره ومواد العروض لأغراض إعداد الصفحات والتسويق والإعلان في قنوات Pakgat وشركائها دون نقل ملكية تلك الحقوق."),
    (7, "السرية وحماية البيانات", "يلتزم الطرفان بسرية المعلومات التجارية والمالية وبيانات العملاء وعدم استخدامها إلا بالقدر اللازم لتنفيذ الاتفاقية أو وفق ما تتطلبه الأنظمة، واتخاذ التدابير المناسبة لحماية البيانات الواقعة تحت سيطرة كل طرف وعدم مشاركتها مع غير المخولين."),
    (8, "حدود الصلاحيات والتعديلات", "لا يعتد بأي تعديل على الشروط التجارية أو الخصومات أو الالتزامات أو الوعود بالمبيعات أو الميزانيات الإعلنية أو الحصرية أو آجال السداد إلا بموافقة كتابية صادرة من Pakgat. ولا يجوز لأي شخص غير مفوض قبض مبالغ باسم Pakgat أو إبرام التزام مالي أو قانوني باسمها."),
    (9, "عدم الحصرية", "ما لم يتفق الطرفان كتابة على خلاف ذلك، تعد هذه الاتفاقية غير حصرية، ويجوز لكل طرف التعامل مع أطراف أخرى شريطة عدم استخدام المعلومات السرية للطرف الآخر أو الإخلال بالطلبات والعروض القائمة وحقوق العملاء."),
    (10, "مدة الاتفاقية وإنهاؤها", "تسري الاتفاقية من تاريخ توقيعها وتستمر حتى انتهاء العروض النشطة وتسوية الالتزامات المالية المتعلقة بها. ويجوز لأي من الطرفين إنهاؤها عند الإخلال الجوهري أو مخالفة الأنظمة أو توقف النشاط أو إساءة استخدام المنصة، مع بقاء حقوق العملاء والمستحقات والالتزامات السابقة نافذة حتى تسويتها."),
    (11, "القوة القاهرة", "لا يتحمل أي من الطرفين مسؤولية التأخير أو عدم التنفيذ الناتج مباشرة عن ظروف خارجة عن الإرادة لا يمكن توقعها أو دفعها بصورة معقولة، مع التزام الطرف المتأثر بإشعار الطرف الآخر واتخاذ ما يمكن لتقليل الأثر."),
    (12, "القانون والاختصاص", "تخضع هذه الاتفاقية لأنظمة المملكة العربية السعودية. ويسعى الطرفان أولاً إلى تسوية أي خلاف ودياً خلال 15 يوم عمل من تاريخ الإشعار الكتابي، فإن تعذر ذلك فيكون الاختصاص للمحاكم المختصة داخل المملكة العربية السعودية."),
    (13, "أحكام عامة", "تمثل هذه الاتفاقية كامل التفاهم بين الطرفين، ولا تعدل إلا كتابة وبموافقتهما. وإذا أصبح أي بند غير نافذ يبقى باقي الاتفاق نافذاً بالقدر الذي يسمح به النظام، ولكل طرف نسخة للعمل بموجبها."),
)


def _logo_bytes() -> bytes:
    path = Path(__file__).resolve().parent / "assets" / "pakgat_contract_reference_logo.b64"
    if not path.is_file():
        raise contract_pdf.ContractRenderError("Pakgat contract logo is missing")
    try:
        payload = base64.b64decode(path.read_text(encoding="ascii").strip(), validate=True)
    except Exception:
        raise contract_pdf.ContractRenderError("Pakgat contract logo is invalid") from None
    if not payload.startswith(b"\x89PNG\r\n\x1a\n"):
        raise contract_pdf.ContractRenderError("Pakgat contract logo is invalid")
    return payload


def _rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi")) or OxmlElement("w:bidi")
    if bidi.getparent() is None:
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return paragraph


def _run(paragraph, text, size, *, bold=False, color=TEXT, rtl=True):
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if rtl:
        run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
    return run


def _cell_text(cell, text, size, *, bold=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    if rtl:
        _rtl(p, align)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    _run(p, text, size, bold=bold, color=color, rtl=rtl)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def _border(cell, color=BORDER, size=4, *, top=True, bottom=True, left=True, right=True):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge, enabled in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if not enabled:
            continue
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)


def _margins(cell, top=40, start=55, bottom=40, end=55):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if mar.getparent() is None:
        tcpr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement("w:" + name)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        mar.append(el)


def _fixed(table):
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    table.autofit = False


def _section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(6)
    section.bottom_margin = Mm(6)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    section.footer_distance = Mm(3)


def _header(doc, logo_path: Path, compact=False):
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed(t)
    for cell, width in zip(t.rows[0].cells, (64, 66, 64)):
        cell.width = Mm(width)
        _margins(cell, 0, 0, 0, 0)
    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=Mm(22 if compact else 25))
    _cell_text(t.cell(0, 2), "اتفاقية شراكة تجارية", 9.5 if compact else 12, bold=True, color=NAVY)
    rule = doc.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = rule.cell(0, 0)
    _cell_text(c, "", 1)
    _border(c, NAVY, 8, top=False, left=False, right=False)


def _footer(section, page, agreement):
    section.footer.is_linked_to_previous = False
    f = section.footer
    f.paragraphs[0].text = ""
    t = f.add_table(rows=1, cols=3, width=Mm(194))
    _fixed(t)
    for c, width in zip(t.rows[0].cells, (64, 66, 64)):
        c.width = Mm(width)
        _border(c, NAVY, 6, bottom=False, left=False, right=False)
        _margins(c, 25, 0, 0, 0)
    _cell_text(t.cell(0, 0), "بكجات | Pakgat.com", 6, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False)
    _cell_text(t.cell(0, 1), agreement, 6, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, rtl=False)
    _cell_text(t.cell(0, 2), f"صفحة {page} من 3", 6, color=MUTED)


def _title(doc, text, size=7.8):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    _cell_text(c, text, size, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _border(c, BORDER, 4, top=False, left=False, right=False)


def _para(doc, text, size, *, bold=False, color=TEXT, center=False, before=0, after=0, keep=False, keep_next=False):
    p = doc.add_paragraph()
    _rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if keep:
        p._p.get_or_add_pPr().append(OxmlElement("w:keepLines"))
    if keep_next:
        p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))
    _run(p, text, size, bold=bold, color=color)


def _parties(doc, data):
    merchant = [
        ("اسم المنشأة", data.legal_name), ("النشاط", data.activity),
        ("السجل التجاري", data.commercial_registration), ("الرقم الضريبي", data.tax_number),
        ("البنك", data.bank_name), ("IBAN", data.iban), ("رقم الجوال", data.contact_phone),
        ("البريد الإلكتروني", data.contact_email), ("الموقع الإلكتروني", data.website or "لا يوجد"),
        ("العنوان", data.national_address), ("اسم الممثل", data.representative_name),
        ("صفة الممثل", data.representative_title),
    ]
    pakgat = [
        ("الاسم", "شركة تام العاصمة التجارية (Pakgat)"), ("السجل التجاري", "1009100740"),
        ("الرقم الضريبي", "312531659100003"), ("الموقع الإلكتروني", "https://pakgat.com"),
        ("IBAN", "SA1710000026700000717001"), ("العنوان", "المملكة العربية السعودية"),
    ] + [("", "")] * 6
    outer = doc.add_table(rows=1, cols=3)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed(outer)
    for c, width in zip(outer.rows[0].cells, (94, 6, 94)):
        c.width = Mm(width)
        _margins(c, 0, 0, 0, 0)

    def card(host, heading, rows):
        t = host.add_table(rows=13, cols=2)
        _fixed(t)
        h = t.rows[0].cells[0]
        h.merge(t.rows[0].cells[1])
        _shade(h, NAVY)
        _cell_text(h, heading, 7.3, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, (label, value) in enumerate(rows, 1):
            v, l = t.rows[i].cells
            _shade(l, LIGHT)
            for c in (v, l):
                _border(c, "E1E7EF", 3)
                _margins(c)
            _cell_text(l, label, 5.7, bold=True, color=NAVY)
            ltr = label in {"IBAN", "رقم الجوال", "البريد الإلكتروني", "الموقع الإلكتروني", "السجل التجاري", "الرقم الضريبي"}
            _cell_text(v, value, 5.55, rtl=not ltr)
    card(outer.cell(0, 0), "الطرف الثاني (التاجر)", merchant)
    card(outer.cell(0, 2), "الطرف الأول", pakgat)


def _signing(doc):
    t = doc.add_table(rows=4, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    lines = (
        ("إجراء التوقيع والاعتماد", 7.6, True),
        ("يراجع التاجر هذه الاتفاقية داخل بوابة Pakgat ويؤكد موافقته عليها باستخدام رمز تحقق OTP مستقل يرسل إلى رقم الجوال المسجل.", 6.7, False),
        ("", 1, False),
        ("نجاح رمز التحقق يعد موافقة إلكترونية موثقة على رقم الاتفاقية المعروض، ثم ينتقل الطلب إلى مراجعة Pakgat النهائية. لا يتم تفعيل حساب التاجر تلقائياً.", 6.4, True),
    )
    for i, (text, size, bold) in enumerate(lines):
        c = t.cell(i, 0)
        _border(c, BORDER, 4, top=i == 0, bottom=i == 3)
        _margins(c, 35, 90, 35, 90)
        _cell_text(c, text, size, bold=bold, color=NAVY if bold else TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)


def _approvals(doc, data):
    left = [
        f"الاسم: {data.representative_name}", f"الصفة: {data.representative_title}",
        f"الجوال: {data.contact_phone}", "الموافقة: إلكترونياً عبر رمز تحقق OTP",
        f"رقم الاتفاقية: {data.agreement_number}",
    ]
    right = [
        "شركة تام العاصمة التجارية (Pakgat)", f"الاسم: {contract_pdf.PAKGAT_SIGNER_NAME}",
        f"الصفة: {contract_pdf.PAKGAT_SIGNER_TITLE}", "الاعتماد: قرار Pakgat النهائي بعد مراجعة الطلب",
        "الحالة: لا يصبح الحساب Active إلا بعد الاعتماد النهائي",
    ]
    outer = doc.add_table(rows=1, cols=3)
    _fixed(outer)
    for c, width in zip(outer.rows[0].cells, (94, 6, 94)):
        c.width = Mm(width)
        _margins(c, 0, 0, 0, 0)
    def card(host, heading, lines):
        t = host.add_table(rows=6, cols=1)
        _fixed(t)
        h = t.cell(0, 0)
        _shade(h, NAVY)
        _cell_text(h, heading, 6.8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, line in enumerate(lines, 1):
            c = t.cell(i, 0)
            _border(c, BORDER, 3, top=False)
            _margins(c)
            _cell_text(c, line, 5.0)
    card(outer.cell(0, 0), "الطرف الثاني (التاجر)", left)
    card(outer.cell(0, 2), "الطرف الأول", right)


def build_contract_docx_otp(data: contract_pdf.ContractData) -> bytes:
    """Build the approved three-page Pakgat contract as a native DOCX."""
    with tempfile.TemporaryDirectory(prefix="pakgat-contract-docx-") as temp:
        root = Path(temp)
        logo = root / "pakgat-logo.png"
        logo.write_bytes(_logo_bytes())
        doc = Document()
        _section(doc.sections[0])
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(7)

        _header(doc, logo)
        _para(doc, "لترويج وبيع العروض والقسائم الإلكترونية بين شركة تام العاصمة التجارية (Pakgat) والتاجر.", 6.8, center=True, after=1.5)
        meta = doc.add_table(rows=1, cols=2)
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, text in zip(meta.rows[0].cells, (f"التاريخ: {data.agreement_date}", f"رقم الاتفاقية: {data.agreement_number}")):
            _shade(c, LIGHT); _border(c); _cell_text(c, text, 7, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        _title(doc, "أولاً: أطراف الاتفاقية")
        _parties(doc, data)
        _title(doc, "ثانياً: التمهيد")
        _para(doc, "حيث إن الطرف الأول يدير منصة إلكترونية متخصصة في تسويق وبيع العروض والباقات والقسائم الإلكترونية، ويرغب الطرف الثاني في عرض خدماته أو منتجاته عبر المنصة للوصول إلى عملاء جدد، فقد اتفق الطرفان – وهما بكامل أهليتهما المعتبرة – على ما يلي، ويعد هذا التمهيد جزءاً لا يتجزأ من الاتفاقية.", 6.25, after=1)
        _signing(doc)
        _footer(doc.sections[0], 1, data.agreement_number)

        second = doc.add_section(WD_SECTION.NEW_PAGE)
        _section(second); _footer(second, 2, data.agreement_number)
        _header(doc, logo, compact=True); _title(doc, "ثالثاً: الشروط والأحكام (1)")
        for n, title, body in CLAUSES[:7]:
            _para(doc, f"{n}. {title}", 6.35, bold=True, color=NAVY, keep=True, keep_next=True)
            _para(doc, body, 5.35, after=1.6, keep=True)

        third = doc.add_section(WD_SECTION.NEW_PAGE)
        _section(third); _footer(third, 3, data.agreement_number)
        _header(doc, logo, compact=True); _title(doc, "ثالثاً: الشروط والأحكام (2)")
        for n, title, body in CLAUSES[7:]:
            _para(doc, f"{n}. {title}", 6.35, bold=True, color=NAVY, keep=True, keep_next=True)
            _para(doc, body, 5.35, after=1.6, keep=True)
        _title(doc, "رابعاً: الموافقة الإلكترونية والاعتماد النهائي", 7.6)
        _approvals(doc, data)
        box = doc.add_table(rows=1, cols=1)
        c = box.cell(0, 0); _shade(c, LIGHT); _border(c); _cell_text(c, ACTIVATION_COPY, 5.7, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

        out = root / "merchant-agreement.docx"
        try:
            doc.save(out)
        except Exception:
            raise contract_pdf.ContractRenderError("Merchant contract DOCX could not be rendered") from None
        return out.read_bytes()


def render_contract_pdf_otp(data: contract_pdf.ContractData, *, converter=contract_pdf._libreoffice_converter) -> bytes:
    with tempfile.TemporaryDirectory(prefix="pakgat-contract-") as temp:
        root = Path(temp)
        source = root / "merchant-agreement.docx"
        source.write_bytes(build_contract_docx_otp(data))
        converter(source, root)
        pdf_path = root / "merchant-agreement.pdf"
        if not pdf_path.exists():
            raise contract_pdf.ContractRenderError("Merchant contract PDF was not generated")
        pdf = pdf_path.read_bytes()
        if not pdf.startswith(b"%PDF"):
            raise contract_pdf.ContractRenderError("Generated merchant contract PDF is invalid")
        return pdf


contract_pdf.build_contract_docx = build_contract_docx_otp
contract_pdf.render_contract_pdf = render_contract_pdf_otp

__all__ = ["ACTIVATION_COPY", "CLAUSES", "build_contract_docx_otp", "render_contract_pdf_otp"]
